## -*- coding: utf-8 -*-
# import pymongo
import docx
# from pymongo import MongoClient

# conn=MongoClient('localhost',27017)
# db=conn['paperIdx']
# paper=db['paper']
# jour=db['jour']
pinfolist=[]
jinfolist=[]
def openFile():
    # r = filedialog.askopenfilename(title='打开文件', filetypes=[('docx', '*.docx'), ('All Files', '*')])
    file = docx.Document('esi.docx')
    for i in range(len(file.paragraphs)):
        infos=file.paragraphs[i].text

        if infos:
            pinfolist.append(infos)

    journalfile=open('journal.txt','r')
    lines=journalfile.readlines()
    for l in lines:
        jinfolist.append(l)
    # print(jinfolist)

def openTableFile():
    journallist=[]
    file = docx.Document('journal.docx')
    t=file.tables[0]
    f=open('journal.txt','a')
    # print(len(t.cols))

    for j in range(1,len(t.rows)):
        temp = []
        for i in range(1,5):
            temp.append(t.cell(j,i).text)

        f.write(','.join(temp))
        f.write('\n')
        f.flush()
    f.close()

def searjInfo(str):
    result = []
    for info in jinfolist:

        if str in info.lower():
            result.append(info)
    return result
def searInfo(str):

    result=[]
    for info in pinfolist:

        if str in info.lower():
            result.append(info)
    return result
openFile()
if __name__=='__main__':
    pass
    # openTableFile()
    # info=searInfo('Some new oscillation results for linear Hamiltonian systems')
    info=searjInfo('7.26')
    print(info)